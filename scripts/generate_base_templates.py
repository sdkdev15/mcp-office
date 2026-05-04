"""Generate base .docx and .pptx templates for all 5 themes."""
import os
import sys
sys.path.insert(0, os.path.dirname(os.path.dirname(os.path.abspath(__file__))))

from docx import Document
from docx.shared import Pt, Inches, Cm, RGBColor
from docx.enum.style import WD_STYLE_TYPE
from pptx import Presentation
from pptx.util import Inches as PptxInches, Pt as PptxPt

from src.styles.themes import THEMES

TEMPLATE_DIR = os.path.join(os.path.dirname(os.path.dirname(os.path.abspath(__file__))), "src", "templates")
os.makedirs(TEMPLATE_DIR, exist_ok=True)


def hex_to_rgb(hex_color: str):
    h = hex_color.lstrip("#")
    return RGBColor(int(h[0:2], 16), int(h[2:4], 16), int(h[4:6], 16))


def generate_docx_template(theme_name: str):
    theme = THEMES[theme_name]
    doc = Document()

    # Page setup
    section = doc.sections[0]
    section.page_width = Cm(21.0)
    section.page_height = Cm(29.7)
    section.top_margin = Cm(2.54)
    section.bottom_margin = Cm(2.54)
    section.left_margin = Cm(2.54)
    section.right_margin = Cm(2.54)

    # Normal style
    normal = doc.styles["Normal"]
    normal.font.name = theme.fonts.body
    normal.font.size = Pt(theme.fonts.body_size)
    normal.font.color.rgb = hex_to_rgb(theme.colors.text)

    # Heading styles
    for level, size in [(1, theme.fonts.heading1_size), (2, theme.fonts.heading2_size),
                         (3, theme.fonts.heading3_size), (4, theme.fonts.heading4_size)]:
        style = doc.styles[f"Heading {level}"]
        style.font.name = theme.fonts.heading
        style.font.size = Pt(size)
        style.font.color.rgb = hex_to_rgb(theme.colors.primary)

    # Title style
    title_style = doc.styles["Title"]
    title_style.font.name = theme.fonts.heading
    title_style.font.size = Pt(28)
    title_style.font.color.rgb = hex_to_rgb(theme.colors.primary)

    # Subtitle style
    subtitle_style = doc.styles["Subtitle"]
    subtitle_style.font.name = theme.fonts.body
    subtitle_style.font.size = Pt(14)
    subtitle_style.font.color.rgb = hex_to_rgb(theme.colors.text_light)
    subtitle_style.font.italic = True

    # Add header with theme label
    header = section.header
    header_para = header.paragraphs[0] if header.paragraphs else header.add_paragraph()
    run = header_para.add_run(theme_name.upper() + " REPORT")
    run.font.size = Pt(8)
    run.font.color.rgb = hex_to_rgb(theme.colors.text_light)
    run.font.name = theme.fonts.body

    # Add a single empty paragraph as placeholder content
    doc.add_paragraph()

    out_path = os.path.join(TEMPLATE_DIR, f"{theme_name}_base.docx")
    doc.save(out_path)
    print(f"  ✅ {out_path} ({os.path.getsize(out_path)} bytes)")


def generate_pptx_template(theme_name: str):
    theme = THEMES[theme_name]
    prs = Presentation()

    # Widescreen 16:9
    prs.slide_width = PptxInches(13.333)
    prs.slide_height = PptxInches(7.5)

    out_path = os.path.join(TEMPLATE_DIR, f"{theme_name}_base.pptx")
    prs.save(out_path)
    print(f"  ✅ {out_path} ({os.path.getsize(out_path)} bytes)")


if __name__ == "__main__":
    print("Generating base templates for all themes...\n")
    for name in THEMES:
        print(f"[{name}]")
        generate_docx_template(name)
        generate_pptx_template(name)
    print(f"\nDone! Generated {len(THEMES) * 2} templates in {TEMPLATE_DIR}")
