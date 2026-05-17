"""Word document generator using python-docx with theme support."""

from __future__ import annotations

import io
import re
from typing import Any, Optional

from docx import Document
from docx.shared import Inches, Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn

from src.styles.themes import get_theme
from src.utils.colors import hex_to_rgbcolor_tuple
from src.utils.metadata import apply_metadata
from src.utils.logger import get_logger
from src.utils.validators import validate_heading_level, ValidationError
from src.utils.image_handler import ImageHandler

log = get_logger("docx_generator")
image_handler = ImageHandler()


def hex_to_rgbcolor(hex_color: str) -> RGBColor:
    """Convert hex color string to RGBColor for python-docx."""
    r, g, b = hex_to_rgbcolor_tuple(hex_color)
    return RGBColor(r, g, b)


class DOCXGenerator:
    """Generates Word documents with rich formatting, tables, and styling."""

    def __init__(self, theme_name: str = "corporate"):
        self.theme_name = theme_name
        self.theme = get_theme(theme_name)
        self.doc: Optional[Document] = None

    def create_document(
        self,
        title: str = "Document",
        page_size: str = "A4",
        orientation: str = "portrait",
        metadata: Optional[dict] = None,
    ) -> bytes:
        """Create a new Word document.

        Args:
            title: Document title.
            page_size: Page size (A4, Letter, Legal).
            orientation: Page orientation (portrait, landscape).
            metadata: Optional document metadata.

        Returns:
            Document content as bytes.
        """
        self.doc = Document()
        self._setup_page(page_size, orientation)
        self._apply_theme_styles()

        # Apply metadata
        if metadata:
            self._apply_metadata(metadata)

        # Add title
        self.add_heading(title, level=1)

        return self._save_document()

    def create_document_with_content(
        self,
        title: str = "Document",
        page_size: str = "A4",
        orientation: str = "portrait",
        metadata: Optional[dict] = None,
        sections: Optional[list[dict]] = None,
        content_paragraphs: Optional[list[str]] = None,
        tables: Optional[list[dict]] = None,
    ) -> bytes:
        """Create a Word document with structured sections.

        Args:
            title: Document title (fallback if no title section provided).
            page_size: Page size (A4, Letter, Legal).
            orientation: Page orientation (portrait, landscape).
            metadata: Optional document metadata.
            sections: Ordered list of section dicts (title, subtitle, toc, heading_1-3, paragraph, list_bullet, list_number, table).
            content_paragraphs: DEPRECATED legacy paragraph list.
            tables: DEPRECATED legacy table list.

        Returns:
            Document content as bytes.
        """
        self.doc = Document()
        self._setup_page(page_size, orientation)
        self._apply_theme_styles()

        if metadata:
            self._apply_metadata(metadata)

        if sections:
            self._process_sections(sections)
        else:
            # Legacy fallback
            self.add_heading(title, level=1)
            if content_paragraphs:
                for text in content_paragraphs:
                    self.add_paragraph(text)
            if tables:
                for table_data in tables:
                    headers = table_data.get("headers", [])
                    rows = table_data.get("rows", [])
                    self.add_table(headers, rows)

        return self._save_document()

    def add_heading(self, text: str, level: int = 1) -> None:
        """Add a heading to the document.

        Args:
            text: Heading text.
            level: Heading level (1-4).
        """
        level = validate_heading_level(level)
        heading = self.doc.add_heading(text, level=level)

        # Apply theme colors
        colors = self.theme.colors
        if level == 1:
            for run in heading.runs:
                pass
            
            if self.theme_name == "corporate":
                try:
                    from docx.shared import Pt, RGBColor
                    header = self.doc.sections[0].header
                    if not header.paragraphs[0].text:
                        hp = header.paragraphs[0]
                        hp.text = "CORPORATE REPORT"
                        hp.style.font.name = "Arial"
                        hp.style.font.size = Pt(9)
                        hp.style.font.color.rgb = RGBColor(156, 163, 175) # Gray-400
                except Exception:
                    pass

    def add_paragraph(
        self,
        text: str,
        bold: bool = False,
        italic: bool = False,
        underline: bool = False,
        strikethrough: bool = False,
        font_name: Optional[str] = None,
        font_size: Optional[int] = None,
        color: Optional[str] = None,
        alignment: Optional[str] = None,
        space_before: Optional[int] = None,
        space_after: Optional[int] = None,
    ) -> None:
        """Add a formatted paragraph.

        Args:
            text: Paragraph text.
            bold: Bold text.
            italic: Italic text.
            underline: Underline text.
            strikethrough: Strikethrough text.
            font_name: Font name.
            font_size: Font size in points.
            color: Hex color string.
            alignment: Text alignment (left, center, right, justify).
            space_before: Space before in points.
            space_after: Space after in points.
        """
        para = self.doc.add_paragraph()
        run = para.add_run(text)

        run.bold = bold
        run.italic = italic
        run.underline = underline
        run.font.strike = strikethrough

        if font_name:
            run.font.name = font_name

        if font_size:
            run.font.size = Pt(font_size)

        if color:
            run.font.color.rgb = hex_to_rgbcolor(color)

        if alignment:
            align_map = {
                "left": WD_ALIGN_PARAGRAPH.LEFT,
                "center": WD_ALIGN_PARAGRAPH.CENTER,
                "right": WD_ALIGN_PARAGRAPH.RIGHT,
                "justify": WD_ALIGN_PARAGRAPH.JUSTIFY,
            }
            para.alignment = align_map.get(alignment.lower(), WD_ALIGN_PARAGRAPH.LEFT)

        if space_before is not None:
            para.paragraph_format.space_before = Pt(space_before)
        if space_after is not None:
            para.paragraph_format.space_after = Pt(space_after)

    def add_table(
        self,
        headers: list[str],
        rows: list[list[Any]],
        style: str = "Table Grid",
    ) -> None:
        """Add a formatted table.

        Args:
            headers: Table headers.
            rows: Table data rows.
            style: Table style name.
        """
        table = self.doc.add_table(rows=1 + len(rows), cols=len(headers))
        table.style = style
        table.alignment = WD_TABLE_ALIGNMENT.CENTER

        colors = self.theme.colors

        # Add headers
        header_row = table.rows[0]
        for i, header in enumerate(headers):
            cell = header_row.cells[i]
            cell.text = str(header)
            for paragraph in cell.paragraphs:
                for run in paragraph.runs:
                    run.bold = True
                    run.font.color.rgb = hex_to_rgbcolor(colors.header_text)
                    run.font.name = self.theme.fonts.body
                paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
            # Set header background
            shading_elm = cell._tc.get_or_add_tcPr().makeelement(qn("w:shd"), {
                qn("w:fill"): colors.header_bg.lstrip("#"),
                qn("w:val"): "clear",
            })
            cell._tc.get_or_add_tcPr().append(shading_elm)

        # Add data rows
        for row_idx, row_data in enumerate(rows):
            table_row = table.rows[row_idx + 1]
            for col_idx, value in enumerate(row_data):
                cell = table_row.cells[col_idx]
                cell.text = str(value)
                for paragraph in cell.paragraphs:
                    for run in paragraph.runs:
                        pass

                # Alternate row shading
                if row_idx % 2 == 1:
                    shading_elm = cell._tc.get_or_add_tcPr().makeelement(qn("w:shd"), {
                        qn("w:fill"): colors.table_alt_row.lstrip("#"),
                        qn("w:val"): "clear",
                    })
                    cell._tc.get_or_add_tcPr().append(shading_elm)

        self.doc.add_paragraph()  # Spacing after table

    def add_list(
        self,
        items: list[str],
        ordered: bool = False,
    ) -> None:
        """Add a bulleted or numbered list.

        Args:
            items: List items.
            ordered: True for numbered list.
        """
        for item in items:
            if ordered:
                self.doc.add_paragraph(item, style="List Number")
            else:
                self.doc.add_paragraph(item, style="List Bullet")

        self.doc.add_paragraph()

    def add_image(
        self,
        image_path: str,
        width: Optional[float] = None,
        height: Optional[float] = None,
        caption: Optional[str] = None,
    ) -> None:
        """Add an image to the document.

        Args:
            image_path: Path to image file.
            width: Width in inches.
            height: Height in inches.
            caption: Optional caption text.
        """
        if width:
            self.doc.add_picture(image_path, width=Inches(width))
        elif height:
            self.doc.add_picture(image_path, height=Inches(height))
        else:
            self.doc.add_picture(image_path, width=Inches(4.0))

        if caption:
            caption_para = self.doc.add_paragraph(caption)
            caption_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
            for run in caption_para.runs:
                run.italic = True
                run.font.size = Pt(10)
                run.font.color.rgb = hex_to_rgbcolor(
                    self.theme.colors.text_light
                )

    def add_page_break(self) -> None:
        """Add a page break."""
        self.doc.add_page_break()

    def add_horizontal_line(self) -> None:
        """Add a horizontal line."""
        p = self.doc.add_paragraph()
        p.paragraph_format.space_before = Pt(6)
        p.paragraph_format.space_after = Pt(6)
        pPr = p._p.get_or_add_pPr()
        pBdr = pPr.makeelement(qn("w:pBdr"), {})
        bottom = pBdr.makeelement(qn("w:bottom"), {
            qn("w:val"): "single",
            qn("w:sz"): "6",
            qn("w:space"): "1",
            qn("w:color"): self.theme.colors.border.lstrip("#"),
        })
        pBdr.append(bottom)
        pPr.append(pBdr)

    def add_header_footer(
        self,
        header_text: Optional[str] = None,
        footer_text: Optional[str] = None,
        include_page_number: bool = True,
    ) -> None:
        """Add header and footer.

        Args:
            header_text: Header text.
            footer_text: Footer text.
            include_page_number: Include page numbers in footer.
        """
        section = self.doc.sections[0]

        if header_text:
            header = section.header
            header_para = header.paragraphs[0]
            header_para.text = header_text
            for run in header_para.runs:
                run.font.size = Pt(9)
                run.font.color.rgb = hex_to_rgbcolor(
                    self.theme.colors.text_light
                )

        if footer_text or include_page_number:
            footer = section.footer
            footer_para = footer.paragraphs[0]

            parts = []
            if footer_text:
                parts.append(footer_text)
            if include_page_number:
                parts.append("Page {PAGE}")

            footer_para.text = " | ".join(parts)
            for run in footer_para.runs:
                run.font.size = Pt(9)

    def add_toc(self) -> None:
        """Add a table of contents placeholder."""
        self.doc.add_paragraph("Table of Contents", style="Heading 2")
        self.doc.add_paragraph("[TOC]", style="TOC 1")
        self.doc.add_page_break()

    def create_from_prompt(self, prompt: str, title: str = "Document") -> bytes:
        """Create a Word document from a natural language prompt.

        Parses the prompt to extract structure: headings, bullet points,
        tables, and paragraphs. Ignores instruction lines (e.g. "Buat dokumen...",
        "Format dokumen dengan...").

        Args:
            prompt: Natural language description with structured content.
            title: Document title.

        Returns:
            Document content as bytes.
        """
        log.info(f"Creating Word document from prompt: {prompt[:100]}...")

        self.doc = Document()
        self._setup_page()
        self._apply_theme_styles()

        # Extract title from prompt if not provided
        doc_title = title if title != "Document" else self._extract_title(prompt)
        self.add_heading(doc_title, level=1)
        self.add_horizontal_line()

        # Parse and render structured content
        self._parse_and_render(prompt)

        return self._save_document()

    def _extract_title(self, text: str) -> str:
        """Extract document title from prompt text."""
        lines = text.strip().split("\n")
        for line in lines:
            line = line.strip()
            # Skip instruction lines
            if re.match(r"^(Buat|Format|Gunakan|Create|Generate)", line, re.IGNORECASE):
                continue
            # First meaningful line is the title
            if line and len(line) > 10:
                # Remove trailing colon
                return line.rstrip(":").strip()
        return "Document"

    def _parse_and_render(self, text: str) -> None:
        """Parse prompt text and render structured content."""
        lines = text.strip().split("\n")
        i = 0
        current_section = None
        bullet_buffer = []
        table_buffer = []
        paragraph_buffer = []

        def flush_bullets():
            nonlocal bullet_buffer
            if bullet_buffer:
                self.add_list(bullet_buffer, ordered=False)
                bullet_buffer = []

        def flush_paragraphs():
            nonlocal paragraph_buffer
            if paragraph_buffer:
                for p in paragraph_buffer:
                    self.add_paragraph(p, space_after=6)
                paragraph_buffer = []

        def flush_table():
            nonlocal table_buffer
            if len(table_buffer) > 1:
                headers = table_buffer[0]
                rows = table_buffer[1:]
                self.add_table(headers, rows)
            table_buffer = []

        while i < len(lines):
            line = lines[i].strip()
            i += 1

            # Skip empty lines
            if not line:
                flush_bullets()
                flush_paragraphs()
                continue

            # Skip instruction lines
            if re.match(r"^(Buat|Format|Gunakan|Create|Generate|Silakan|Please)", line, re.IGNORECASE):
                flush_bullets()
                flush_paragraphs()
                continue

            # Detect section heading patterns:
            # 1. ALL CAPS
            # 2. Ends with colon
            # 3. "Heading - description" pattern (title case word followed by dash)
            is_heading = False
            heading_text = None
            dash_match = None

            if line.isupper() and not line.startswith(("-", "*", "•")):
                is_heading = True
                heading_text = line.title()
            elif line.endswith(":") and not line.startswith(("-", "*", "•")):
                is_heading = True
                heading_text = line.rstrip(":").strip()
            elif not line.startswith(("-", "*", "•", "1.", "2.", "3.")):
                # Check for "Heading - content" pattern (lenient)
                dash_match = re.match(r"^([A-Z][^-\n]{5,}?)\s+-\s+(.+)$", line)
                if dash_match:
                    is_heading = True
                    heading_text = dash_match.group(1).strip()

            if is_heading and heading_text:
                flush_bullets()
                flush_paragraphs()
                flush_table()
                self.add_heading(heading_text, level=2)
                current_section = heading_text
                # If there was content after the dash, add it as a paragraph
                if dash_match and dash_match.group(2):
                    self.add_paragraph(dash_match.group(2).strip(), space_after=6)
                continue

            # Detect bullet points
            bullet_match = re.match(r"^[-*•]\s+(.*)", line)
            if bullet_match:
                flush_paragraphs()
                flush_table()
                bullet_buffer.append(bullet_match.group(1).strip())
                continue

            # Detect numbered items
            num_match = re.match(r"^\d+\.\s+(.*)", line)
            if num_match:
                flush_paragraphs()
                flush_table()
                bullet_buffer.append(num_match.group(1).strip())
                continue

            # Detect table-like content (pipe-separated or tab-separated)
            if "|" in line and line.count("|") > 1:
                flush_bullets()
                flush_paragraphs()
                cells = [c.strip() for c in line.split("|") if c.strip()]
                if cells:
                    table_buffer.append(cells)
                continue

            # Regular paragraph
            flush_bullets()
            flush_table()
            paragraph_buffer.append(line)

        # Flush remaining buffers
        flush_bullets()
        flush_paragraphs()
        flush_table()

    def _setup_page(self, page_size: str = "A4", orientation: str = "portrait") -> None:
        """Setup page size and orientation."""
        section = self.doc.sections[0]

        if page_size.upper() == "A4":
            section.page_width = Cm(21.0)
            section.page_height = Cm(29.7)
        elif page_size.upper() == "LETTER":
            section.page_width = Inches(8.5)
            section.page_height = Inches(11)
        elif page_size.upper() == "LEGAL":
            section.page_width = Inches(8.5)
            section.page_height = Inches(14)

        if orientation.lower() == "landscape":
            section.orientation = 1  # LANDSCAPE

        # Set margins
        section.top_margin = Cm(2.54)
        section.bottom_margin = Cm(2.54)
        section.left_margin = Cm(2.54)
        section.right_margin = Cm(2.54)

    def _apply_theme_styles(self) -> None:
        """Apply theme styles to the document style definitions."""
        style = self.doc.styles["Normal"]
        font = style.font
        font.name = self.theme.fonts.body
        font.size = Pt(self.theme.fonts.body_size)

    def add_table_of_contents(self, title: str = "Table of Contents") -> None:
        """Insert a Word-native Table of Contents field.

        The TOC uses field codes that Word/WPS Office will render and
        auto-populate from Heading 1-3 styles when the user updates fields.
        """
        from docx.oxml import OxmlElement

        # Add TOC heading
        toc_heading = self.doc.add_paragraph(title)
        toc_heading.style = self.doc.styles['Heading 1']

        # Create the TOC field using complex field characters
        paragraph = self.doc.add_paragraph()
        run = paragraph.add_run()
        fldChar_begin = OxmlElement('w:fldChar')
        fldChar_begin.set(qn('w:fldCharType'), 'begin')
        run._r.append(fldChar_begin)

        run2 = paragraph.add_run()
        instrText = OxmlElement('w:instrText')
        instrText.set(qn('xml:space'), 'preserve')
        instrText.text = ' TOC \\o "1-3" \\h \\z \\u '
        run2._r.append(instrText)

        run3 = paragraph.add_run()
        fldChar_separate = OxmlElement('w:fldChar')
        fldChar_separate.set(qn('w:fldCharType'), 'separate')
        run3._r.append(fldChar_separate)

        # Placeholder text shown before field update
        run4 = paragraph.add_run('[Update this field to populate Table of Contents]')
        run4.italic = True
        run4.font.color.rgb = RGBColor(156, 163, 175)

        run5 = paragraph.add_run()
        fldChar_end = OxmlElement('w:fldChar')
        fldChar_end.set(qn('w:fldCharType'), 'end')
        run5._r.append(fldChar_end)

        # Add a page break after TOC
        self.doc.add_page_break()

    def add_subtitle(self, text: str) -> None:
        """Add a subtitle paragraph using the built-in 'Subtitle' style."""
        self.doc.add_paragraph(text, style='Subtitle')

    def _process_sections(self, sections: list[dict]) -> None:
        """Process an ordered array of document sections.

        Each section dict must have a 'type' key. Supported types:
        title, subtitle, toc, heading_1, heading_2, heading_3,
        paragraph, list_bullet, list_number, table.
        """
        for section in sections:
            section_type = section.get("type", "paragraph")
            text = section.get("text", "")

            if section_type == "title":
                self.doc.add_paragraph(text, style='Title')
            elif section_type == "subtitle":
                self.add_subtitle(text)
            elif section_type == "toc":
                self.add_table_of_contents(section.get("title", "Table of Contents"))
            elif section_type == "heading_1":
                self.add_heading(text, level=1)
            elif section_type == "heading_2":
                self.add_heading(text, level=2)
            elif section_type == "heading_3":
                self.add_heading(text, level=3)
            elif section_type == "paragraph":
                self.add_paragraph(text)
            elif section_type == "list_bullet":
                items = section.get("items", [])
                self.add_list(items, ordered=False)
            elif section_type == "list_number":
                items = section.get("items", [])
                self.add_list(items, ordered=True)
            elif section_type == "table":
                headers = section.get("headers", [])
                rows = section.get("rows", [])
                if headers and rows:
                    self.add_table(headers, rows)
            elif section_type == "image":
                src = section.get("source")
                width = section.get("width")
                height = section.get("height")
                caption = section.get("caption")
                if src:
                    try:
                        # docx add_picture accepts local paths or file-like objects
                        cached_path = image_handler.process_image(src)
                        self.add_image(cached_path, width=width, height=height, caption=caption)
                    except Exception as e:
                        log.warning(f"Failed to add image to docx: {e}")
            else:
                log.warning(f"Unknown section type '{section_type}', treating as paragraph")
                if text:
                    self.add_paragraph(text)

    def _apply_metadata(self, metadata: dict) -> None:
        """Apply document metadata."""
        apply_metadata(self.doc.core_properties, metadata)

    def create_from_template(
        self,
        template_path: str,
        title: str = "Document",
        sections: Optional[list[dict]] = None,
        content_paragraphs: Optional[list[str]] = None,
        tables: Optional[list[dict]] = None,
        metadata: Optional[dict] = None,
    ) -> bytes:
        """Create a document using an existing .docx file as template.

        The template preserves all formatting, styles, headers, footers,
        and page setup. New content is appended after existing content.

        Args:
            template_path: Path to the .docx template file.
            title: Document title (replaces first Heading 1 if found).
            sections: Ordered list of section dicts.
            content_paragraphs: DEPRECATED legacy paragraph list.
            tables: DEPRECATED legacy table list.
            metadata: Optional document metadata.

        Returns:
            Document content as bytes.
        """
        self.doc = Document(template_path)

        # Apply metadata
        if metadata:
            self._apply_metadata(metadata)

        if sections:
            # New structured sections flow
            self._process_sections(sections)
        else:
            # Legacy fallback
            # Replace first heading with title if found
            if title:
                for para in self.doc.paragraphs:
                    if para.style and para.style.name and para.style.name.startswith("Heading"):
                        if para.runs:
                            para.runs[0].text = title
                            for run in para.runs[1:]:
                                run.text = ""
                        else:
                            para.text = title
                        break

            if content_paragraphs:
                for text in content_paragraphs:
                    self.add_paragraph(text)

            if tables:
                for table_data in tables:
                    headers = table_data.get("headers", [])
                    rows = table_data.get("rows", [])
                    self.add_table(headers, rows)

        return self._save_document()

    def _save_document(self) -> bytes:
        """Save document to bytes."""
        buffer = io.BytesIO()
        self.doc.save(buffer)
        buffer.seek(0)
        return buffer.getvalue()
